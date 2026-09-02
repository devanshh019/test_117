# Local RAG Vector Knowledge Base with ChromaDB and Embeddings
import time
from pathlib import Path
from typing import Dict, List, Any, Optional

import docx
from chromadb.utils import embedding_functions
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document

from .config import (
    CHROMA_DIR,
    SEED_DOCS_DIR,
    KB_DOCS_DIR,
    RAG_CHUNK_SIZE,
    RAG_CHUNK_OVERLAP,
    RAG_DEFAULT_TOP_K,
)


# Local ONNX Embedding Function & Text Splitter
emb_fn = embedding_functions.DefaultEmbeddingFunction()


class LocalEmbeddings:
    def embed_documents(self, texts):
        return emb_fn(texts)

    def embed_query(self, text):
        return emb_fn([text])[0]


splitter = RecursiveCharacterTextSplitter(
    chunk_size=RAG_CHUNK_SIZE,
    chunk_overlap=RAG_CHUNK_OVERLAP,
)

vec_store = Chroma(
    collection_name="kavach_standards",
    embedding_function=LocalEmbeddings(),
    persist_directory=str(CHROMA_DIR),
)


# Simple Knowledge Base Interface (< 100 lines)
class LocalRAGKnowledgeBase:

    @staticmethod
    def extract_text_from_file(file_path: Path or str) -> str:
        """Robust multi-format text extractor for PDF, DOCX, DOC, TXT, CSV, MD files."""
        p = Path(file_path)
        ext = p.suffix.lower()

        if ext in [".docx", ".doc"]:
            try:
                doc = docx.Document(str(p))
                paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        row_txt = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_txt:
                            paragraphs.append(" | ".join(row_txt))
                return "\n\n".join(paragraphs)
            except Exception as e:
                return f"[DOCX Extraction Error: {e}]"
        elif ext == ".pdf":
            try:
                loader = PyPDFLoader(str(p))
                pages = loader.load()
                return "\n\n".join(page.page_content for page in pages)
            except Exception as e:
                return f"[PDF Extraction Error: {e}]"
        else:
            try:
                with open(p, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            except Exception as e:
                return f"[Text Read Error: {e}]"

    def ingest_file(self, file_path: str or Path, original_filename: str = None) -> Dict[str, Any]:
        """Ingests and indexes PDF, DOCX, DOC, TXT, CSV, and Markdown files into local Chroma vector store."""
        p = Path(file_path)
        name = original_filename or p.name
        doc_id = f"DOC-{int(time.time() * 1000)}"
        ext = p.suffix.lower()

        if ext == ".pdf":
            try:
                docs = PyPDFLoader(str(p)).load()
                for d in docs:
                    d.metadata["doc_id"] = doc_id
                    d.metadata["title"] = name
                    d.metadata["filename"] = name
            except Exception:
                raw_text = self.extract_text_from_file(p)
                docs = [Document(page_content=raw_text, metadata={"doc_id": doc_id, "title": name, "filename": name})]
        else:
            raw_text = self.extract_text_from_file(p)
            if not raw_text or raw_text.startswith("["):
                raise ValueError(f"Could not extract readable text from '{name}'.")
            docs = [Document(page_content=raw_text, metadata={"doc_id": doc_id, "title": name, "filename": name})]

        chunks = splitter.split_documents(docs)
        if not chunks:
            raise ValueError(f"Document '{name}' contains no indexable text.")

        vec_store.add_documents(chunks)
        return {
            "success": True,
            "filename": name,
            "doc_id": doc_id,
            "indexed_chunks": len(chunks),
            "document": {"doc_id": doc_id, "filename": name, "title": name},
        }

    def ingest_text(self, title: str, text: str, category: str = "STANDARDS") -> Dict[str, Any]:
        doc_id = f"DOC-{int(time.time() * 1000)}"
        doc = Document(page_content=text, metadata={"doc_id": doc_id, "title": title, "filename": title})
        chunks = splitter.split_documents([doc])
        vec_store.add_documents(chunks)
        return {
            "success": True,
            "filename": title,
            "doc_id": doc_id,
            "indexed_chunks": len(chunks),
            "document": {"doc_id": doc_id, "title": title},
        }

    def search(self, query: str, top_k: int = RAG_DEFAULT_TOP_K) -> List[Dict[str, Any]]:
        results = vec_store.similarity_search_with_score(query, k=top_k)
        citations = []
        for doc, dist in results:
            if dist > 1.72:
                continue

            relevance = max(0.5, min(0.99, round(1.0 - (dist / 3.0), 2)))
            citations.append({
                "doc_id": doc.metadata.get("doc_id", "DOC"),
                "title": doc.metadata.get("title", "Standard Document"),
                "filename": doc.metadata.get("filename", "document.txt"),
                "chunk_index": 1,
                "total_chunks": 1,
                "excerpt": doc.page_content[:250] + "...",
                "full_content": doc.page_content,
                "relevance_score": relevance,
                "distance": round(dist, 3)
            })
        return citations

    def list_documents(self) -> List[Dict[str, Any]]:
        metas = vec_store.get(include=["metadatas"]).get("metadatas", [])
        seen = {}
        for m in metas:
            if m.get("doc_id") and m["doc_id"] not in seen:
                seen[m["doc_id"]] = {
                    "doc_id": m["doc_id"],
                    "title": m.get("title", "Document"),
                    "filename": m.get("filename", "document.txt"),
                }
        return list(seen.values())

    def delete_document(self, doc_id: str) -> bool:
        ids = vec_store.get(where={"doc_id": doc_id}).get("ids", [])
        if ids:
            vec_store.delete(ids=ids)
            return True
        return False

    def get_stats(self) -> Dict[str, Any]:
        ids = vec_store.get().get("ids", [])
        return {"total_documents": len(self.list_documents()), "total_chunks": len(ids)}


# Ingest Seed Documents on Startup
def _seed_db():
    if SEED_DOCS_DIR.exists() and len(vec_store.get().get("ids", [])) == 0:
        for p in SEED_DOCS_DIR.glob("*.*"):
            try:
                knowledge_base.ingest_file(p, original_filename=p.name)
            except Exception:
                pass


knowledge_base = LocalRAGKnowledgeBase()
_seed_db()
