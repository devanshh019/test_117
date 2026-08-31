import os
os.environ['MPLCONFIGDIR'] = '/Users/devanshkumarverma/Desktop/PS117-v2.0.0/backend/data/storage/on_premises_cache'
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

_target_plot_path = r'/Users/devanshkumarverma/Desktop/PS117-v2.0.0/backend/data/storage/plot_exec_sim_1788158251606.png'
_orig_savefig = plt.savefig
def _patched_savefig(*args, **kwargs):
    kwargs['dpi'] = kwargs.get('dpi', 300)
    kwargs['bbox_inches'] = kwargs.get('bbox_inches', 'tight')
    _orig_savefig(_target_plot_path, **kwargs)
plt.savefig = _patched_savefig

_orig_close = plt.close
def _patched_close(*args, **kwargs):
    if plt.get_fignums() and not os.path.exists(_target_plot_path):
        try:
            _orig_savefig(_target_plot_path, dpi=300, bbox_inches='tight')
        except Exception:
            pass
    return _orig_close(*args, **kwargs)
plt.close = _patched_close

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('PSU Approval Note - Distillation Column C-101', 0)

document.add_paragraph('**Date:** 2024-01-26')
document.add_paragraph('**Subject:** UTT Review & Remaining Life Assessment')
document.add_paragraph('**Column:** C-101')
document.add_paragraph('**Operating Conditions:** 350°C, 18.5 bar')

document.add_paragraph('**Summary:**')
document.add_paragraph('This document summarizes the review of the Ultrasonic Thickness Testing (UTT) report for Distillation Column C-101. Based on the provided data and assumptions, the calculated corrosion rate is 1.0 mm/year, and the remaining life is 10 years.  A full ASME Sec VIII code check is required to determine the actual wall thickness and ensure compliance with design specifications.')

document.add_paragraph('**Recommendations:**')
document.add_paragraph('1. Conduct a full ASME Sec VIII code check, incorporating the calculated corrosion rate and the actual design pressure.')
document.add_paragraph('2.  Verify the fluid composition to accurately determine the corrosion rate.')
document.add_paragraph('3.  Schedule a turnaround if the remaining life is less than 2 years.')

document.add_paragraph('**Prepared By:** KAVACH-AI')

document.save('PSU_Approval_Note_C101.docx')

if plt.get_fignums() and not os.path.exists(_target_plot_path):
    try:
        _orig_savefig(_target_plot_path, dpi=300, bbox_inches='tight')
    except Exception:
        pass
