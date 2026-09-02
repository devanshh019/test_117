# Word Document (.docx) Strategy Renderer
from __future__ import annotations

import os
from pathlib import Path

import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .base import Renderer
from ..schemas import DocxSpec, DocumentResult, DocType


class DocxRenderer(Renderer):
    """Port the body of the old generate_custom_word_doc() here almost
    unchanged — the win isn't rewriting the docx-building logic, it's that
    spec is now a validated DocxSpec instead of a raw dict with .get()
    everywhere."""

    def render(self, spec: DocxSpec, output_path: Path) -> DocumentResult:
        doc = docx.Document()
        style = doc.styles["Normal"]
        style.font.name = self.theme.font_family
        style.font.color.rgb = RGBColor(*self.theme.rgb_dark)

        p_hdr = doc.add_paragraph()
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_hdr.add_run(spec.title.upper())
        run_title.bold = True
        run_title.font.size = Pt(self.theme.header_size)

        if spec.subject:
            p_subj = doc.add_paragraph()
            r = p_subj.add_run("SUBJECT: ")
            r.bold = True
            r.font.color.rgb = RGBColor(*self.theme.rgb_orange)
            p_subj.add_run(spec.subject)

        for section in spec.sections:
            doc.add_heading(section.heading, level=2)
            for chunk in section.content.split("\n\n"):
                if chunk.strip():
                    doc.add_paragraph(chunk.strip())

        doc.save(output_path)

        return DocumentResult(
            filename=output_path.name,
            path=str(output_path),
            doc_type=DocType.DOCX,
            size_bytes=os.path.getsize(output_path),
        )
